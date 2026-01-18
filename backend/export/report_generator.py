from typing import Dict, List, Any, Optional
import structlog
from datetime import datetime
import io
import base64

logger = structlog.get_logger()

class ReportGenerator:
    """
    Generates reports in multiple formats from orchestration results.
    """
    
    def __init__(self):
        self.supported_formats = [
            "pdf", "docx", "xlsx", "csv", "png", "jpeg", "json", "html", "markdown"
        ]
    
    def generate_report(
        self,
        data: Dict[str, Any],
        format: str,
        template: Optional[str] = None
    ) -> bytes:
        """
        Generate report in specified format.
        """
        format = format.lower()
        
        if format not in self.supported_formats:
            raise ValueError(f"Unsupported format: {format}")
        
        logger.info("generating_report", format=format, template=template)
        
        if format == "pdf":
            return self._generate_pdf(data)
        elif format == "docx":
            return self._generate_docx(data)
        elif format == "xlsx":
            return self._generate_excel(data)
        elif format == "csv":
            return self._generate_csv(data)
        elif format == "png":
            return self._generate_image(data, "png")
        elif format == "jpeg":
            return self._generate_image(data, "jpeg")
        elif format == "json":
            return self._generate_json(data)
        elif format == "html":
            return self._generate_html(data)
        elif format == "markdown":
            return self._generate_markdown(data)
    
    def _generate_pdf(self, data: Dict) -> bytes:
        """
        Generate PDF report using reportlab.
        """
        try:
            from reportlab.lib.pagesizes import letter, A4
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
            from reportlab.lib import colors
            from reportlab.lib.enums import TA_CENTER, TA_LEFT
            
            buffer = io.BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=letter)
            story = []
            styles = getSampleStyleSheet()
            
            # Custom styles
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=24,
                textColor=colors.HexColor('#667eea'),
                spaceAfter=30,
                alignment=TA_CENTER
            )
            
            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading2'],
                fontSize=16,
                textColor=colors.HexColor('#333333'),
                spaceAfter=12
            )
            
            # Title
            story.append(Paragraph("Network Consultant AI Report", title_style))
            story.append(Spacer(1, 0.2*inch))
            
            # Metadata
            story.append(Paragraph("Report Details", heading_style))
            metadata_data = [
                ['Generated:', datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S UTC')],
                ['Request ID:', data.get('request_id', 'N/A')],
                ['Priority:', data.get('priority', 'N/A')],
                ['Confidence:', f"{data.get('confidence', 0) * 100:.1f}%"]
            ]
            
            metadata_table = Table(metadata_data, colWidths=[2*inch, 4*inch])
            metadata_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#f0f0f0')),
                ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 10),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey)
            ]))
            story.append(metadata_table)
            story.append(Spacer(1, 0.3*inch))
            
            # Issue
            story.append(Paragraph("Client Issue", heading_style))
            story.append(Paragraph(data.get('client_issue', 'N/A'), styles['Normal']))
            story.append(Spacer(1, 0.2*inch))
            
            # Consensus
            story.append(Paragraph("Consensus Analysis", heading_style))
            story.append(Paragraph(data.get('consensus', 'N/A'), styles['Normal']))
            story.append(Spacer(1, 0.2*inch))
            
            # Recommendations
            if 'recommendations' in data and data['recommendations']:
                story.append(Paragraph("Recommendations", heading_style))
                for i, rec in enumerate(data['recommendations'], 1):
                    story.append(Paragraph(f"{i}. {rec}", styles['Normal']))
                    story.append(Spacer(1, 0.1*inch))
            
            # Agent Analysis
            if 'agents' in data:
                story.append(PageBreak())
                story.append(Paragraph("Agent Analysis", heading_style))
                
                for agent_name, agent_data in data['agents'].items():
                    story.append(Paragraph(f"<b>{agent_name}</b>", styles['Heading3']))
                    story.append(Paragraph(agent_data.get('analysis', 'N/A'), styles['Normal']))
                    story.append(Spacer(1, 0.15*inch))
            
            doc.build(story)
            buffer.seek(0)
            return buffer.getvalue()
            
        except ImportError:
            logger.error("reportlab_not_installed")
            raise ImportError("Install reportlab: pip install reportlab")
    
    def _generate_docx(self, data: Dict) -> bytes:
        """
        Generate DOCX report using python-docx.
        """
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            doc = Document()
            
            # Title
            title = doc.add_heading('Network Consultant AI Report', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title.runs[0].font.color.rgb = RGBColor(102, 126, 234)
            
            # Metadata
            doc.add_heading('Report Details', level=1)
            table = doc.add_table(rows=4, cols=2)
            table.style = 'Light Grid Accent 1'
            
            cells = [
                ['Generated:', datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S UTC')],
                ['Request ID:', data.get('request_id', 'N/A')],
                ['Priority:', data.get('priority', 'N/A')],
                ['Confidence:', f"{data.get('confidence', 0) * 100:.1f}%"]
            ]
            
            for i, (key, value) in enumerate(cells):
                table.rows[i].cells[0].text = key
                table.rows[i].cells[1].text = value
            
            # Issue
            doc.add_heading('Client Issue', level=1)
            doc.add_paragraph(data.get('client_issue', 'N/A'))
            
            # Consensus
            doc.add_heading('Consensus Analysis', level=1)
            doc.add_paragraph(data.get('consensus', 'N/A'))
            
            # Recommendations
            if 'recommendations' in data and data['recommendations']:
                doc.add_heading('Recommendations', level=1)
                for rec in data['recommendations']:
                    doc.add_paragraph(rec, style='List Bullet')
            
            # Agent Analysis
            if 'agents' in data:
                doc.add_page_break()
                doc.add_heading('Agent Analysis', level=1)
                
                for agent_name, agent_data in data['agents'].items():
                    doc.add_heading(agent_name, level=2)
                    doc.add_paragraph(agent_data.get('analysis', 'N/A'))
            
            # Red Flag Warning
            if data.get('red_flagged'):
                doc.add_paragraph()
                warning = doc.add_paragraph('⚠️ RED FLAG: This issue requires immediate attention!')
                warning.runs[0].font.color.rgb = RGBColor(220, 53, 69)
                warning.runs[0].font.bold = True
            
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer.getvalue()
            
        except ImportError:
            logger.error("python_docx_not_installed")
            raise ImportError("Install python-docx: pip install python-docx")
    
    def _generate_excel(self, data: Dict) -> bytes:
        """
        Generate Excel report using openpyxl.
        """
        try:
            from openpyxl import Workbook
            from openpyxl.styles import Font, PatternFill, Alignment
            
            wb = Workbook()
            
            # Summary Sheet
            ws_summary = wb.active
            ws_summary.title = "Summary"
            
            # Headers
            ws_summary['A1'] = "Network Consultant AI Report"
            ws_summary['A1'].font = Font(size=16, bold=True, color="667eea")
            ws_summary.merge_cells('A1:B1')
            
            # Metadata
            row = 3
            ws_summary[f'A{row}'] = "Generated:"
            ws_summary[f'B{row}'] = datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S UTC')
            row += 1
            
            ws_summary[f'A{row}'] = "Request ID:"
            ws_summary[f'B{row}'] = data.get('request_id', 'N/A')
            row += 1
            
            ws_summary[f'A{row}'] = "Priority:"
            ws_summary[f'B{row}'] = data.get('priority', 'N/A')
            row += 1
            
            ws_summary[f'A{row}'] = "Confidence:"
            ws_summary[f'B{row}'] = f"{data.get('confidence', 0) * 100:.1f}%"
            row += 2
            
            # Issue
            ws_summary[f'A{row}'] = "Client Issue:"
            ws_summary[f'A{row}'].font = Font(bold=True)
            row += 1
            ws_summary[f'A{row}'] = data.get('client_issue', 'N/A')
            ws_summary.merge_cells(f'A{row}:B{row}')
            row += 2
            
            # Consensus
            ws_summary[f'A{row}'] = "Consensus:"
            ws_summary[f'A{row}'].font = Font(bold=True)
            row += 1
            ws_summary[f'A{row}'] = data.get('consensus', 'N/A')
            ws_summary.merge_cells(f'A{row}:B{row}')
            
            # Recommendations Sheet
            if 'recommendations' in data and data['recommendations']:
                ws_rec = wb.create_sheet("Recommendations")
                ws_rec['A1'] = "Recommendations"
                ws_rec['A1'].font = Font(size=14, bold=True)
                
                for i, rec in enumerate(data['recommendations'], 2):
                    ws_rec[f'A{i}'] = rec
            
            # Agent Analysis Sheet
            if 'agents' in data:
                ws_agents = wb.create_sheet("Agent Analysis")
                ws_agents['A1'] = "Agent"
                ws_agents['B1'] = "Analysis"
                ws_agents['A1'].font = Font(bold=True)
                ws_agents['B1'].font = Font(bold=True)
                
                row = 2
                for agent_name, agent_data in data['agents'].items():
                    ws_agents[f'A{row}'] = agent_name
                    ws_agents[f'B{row}'] = agent_data.get('analysis', 'N/A')
                    row += 1
            
            # Auto-size columns
            for ws in wb.worksheets:
                for column in ws.columns:
                    max_length = 0
                    column = [cell for cell in column]
                    for cell in column:
                        try:
                            if len(str(cell.value)) > max_length:
                                max_length = len(cell.value)
                        except:
                            pass
                    adjusted_width = min(max_length + 2, 100)
                    ws.column_dimensions[column[0].column_letter].width = adjusted_width
            
            buffer = io.BytesIO()
            wb.save(buffer)
            buffer.seek(0)
            return buffer.getvalue()
            
        except ImportError:
            logger.error("openpyxl_not_installed")
            raise ImportError("Install openpyxl: pip install openpyxl")
    
    def _generate_csv(self, data: Dict) -> bytes:
        """
        Generate CSV export.
        """
        import csv
        
        buffer = io.StringIO()
        writer = csv.writer(buffer)
        
        # Headers
        writer.writerow(['Field', 'Value'])
        
        # Metadata
        writer.writerow(['Generated', datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S UTC')])
        writer.writerow(['Request ID', data.get('request_id', 'N/A')])
        writer.writerow(['Priority', data.get('priority', 'N/A')])
        writer.writerow(['Confidence', f"{data.get('confidence', 0) * 100:.1f}%"])
        writer.writerow(['Red Flagged', data.get('red_flagged', False)])
        writer.writerow([])
        
        # Issue
        writer.writerow(['Client Issue', data.get('client_issue', 'N/A')])
        writer.writerow([])
        
        # Consensus
        writer.writerow(['Consensus', data.get('consensus', 'N/A')])
        writer.writerow([])
        
        # Recommendations
        if 'recommendations' in data:
            writer.writerow(['Recommendations'])
            for rec in data['recommendations']:
                writer.writerow(['', rec])
            writer.writerow([])
        
        # Agents
        if 'agents' in data:
            writer.writerow(['Agent', 'Analysis'])
            for agent_name, agent_data in data['agents'].items():
                writer.writerow([agent_name, agent_data.get('analysis', 'N/A')])
        
        content = buffer.getvalue()
        return content.encode('utf-8')
    
    def _generate_image(self, data: Dict, format: str) -> bytes:
        """
        Generate image report (PNG/JPEG) using Pillow.
        """
        try:
            from PIL import Image, ImageDraw, ImageFont
            
            # Create image
            width, height = 800, 1200
            bg_color = (255, 255, 255)
            img = Image.new('RGB', (width, height), bg_color)
            draw = ImageDraw.Draw(img)
            
            # Try to use a nice font, fallback to default
            try:
                title_font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", 24)
                heading_font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", 16)
                normal_font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 12)
            except:
                title_font = ImageFont.load_default()
                heading_font = ImageFont.load_default()
                normal_font = ImageFont.load_default()
            
            y_position = 30
            padding = 20
            
            # Title
            draw.text((padding, y_position), "Network Consultant AI Report", 
                     fill=(102, 126, 234), font=title_font)
            y_position += 50
            
            # Metadata
            draw.text((padding, y_position), "Report Details", fill=(51, 51, 51), font=heading_font)
            y_position += 30
            
            metadata = [
                f"Generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S UTC')}",
                f"Request ID: {data.get('request_id', 'N/A')}",
                f"Priority: {data.get('priority', 'N/A')}",
                f"Confidence: {data.get('confidence', 0) * 100:.1f}%"
            ]
            
            for line in metadata:
                draw.text((padding, y_position), line, fill=(0, 0, 0), font=normal_font)
                y_position += 20
            
            y_position += 20
            
            # Issue
            draw.text((padding, y_position), "Client Issue", fill=(51, 51, 51), font=heading_font)
            y_position += 30
            
            issue_text = data.get('client_issue', 'N/A')
            words = issue_text.split()
            line = ""
            for word in words:
                test_line = line + word + " "
                if len(test_line) * 7 < width - (padding * 2):
                    line = test_line
                else:
                    draw.text((padding, y_position), line, fill=(0, 0, 0), font=normal_font)
                    y_position += 20
                    line = word + " "
            if line:
                draw.text((padding, y_position), line, fill=(0, 0, 0), font=normal_font)
                y_position += 20
            
            y_position += 20
            
            # Red Flag Warning
            if data.get('red_flagged'):
                draw.text((padding, y_position), "⚠️ RED FLAG: Immediate attention required!", 
                         fill=(220, 53, 69), font=heading_font)
            
            buffer = io.BytesIO()
            img.save(buffer, format=format.upper())
            buffer.seek(0)
            return buffer.getvalue()
            
        except ImportError:
            logger.error("pillow_not_installed")
            raise ImportError("Install Pillow: pip install Pillow")
    
    def _generate_json(self, data: Dict) -> bytes:
        import json
        return json.dumps(data, indent=2).encode('utf-8')
    
    def _generate_html(self, data: Dict) -> bytes:
        html = f"""
<!DOCTYPE html>
<html>
<head>
    <title>Network Consultant AI Report</title>
    <style>
        body {{ font-family: Arial, sans-serif; max-width: 900px; margin: 40px auto; padding: 20px; }}
        h1 {{ color: #667eea; text-align: center; }}
        h2 {{ color: #333; border-bottom: 2px solid #667eea; padding-bottom: 10px; }}
        .metadata {{ background: #f5f7fa; padding: 15px; border-radius: 8px; margin: 20px 0; }}
        .warning {{ background: #ffebee; color: #c62828; padding: 15px; border-left: 4px solid #c62828; margin: 20px 0; }}
        .recommendation {{ background: #e8f5e9; padding: 10px; margin: 10px 0; border-radius: 4px; }}
    </style>
</head>
<body>
    <h1>Network Consultant AI Report</h1>
    
    <div class="metadata">
        <p><strong>Generated:</strong> {datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S UTC')}</p>
        <p><strong>Request ID:</strong> {data.get('request_id', 'N/A')}</p>
        <p><strong>Priority:</strong> {data.get('priority', 'N/A')}</p>
        <p><strong>Confidence:</strong> {data.get('confidence', 0) * 100:.1f}%</p>
    </div>
    
    {'<div class="warning">⚠️ RED FLAG: This issue requires immediate attention!</div>' if data.get('red_flagged') else ''}
    
    <h2>Client Issue</h2>
    <p>{data.get('client_issue', 'N/A')}</p>
    
    <h2>Consensus Analysis</h2>
    <p>{data.get('consensus', 'N/A')}</p>
    
    <h2>Recommendations</h2>
    {''.join(f'<div class="recommendation">{rec}</div>' for rec in data.get('recommendations', []))}
    
    <h2>Agent Analysis</h2>
    {''.join(f'<h3>{name}</h3><p>{agent.get("analysis", "N/A")}</p>' for name, agent in data.get('agents', {}).items())}
</body>
</html>
        """
        return html.encode('utf-8')
    
    def _generate_markdown(self, data: Dict) -> bytes:
        md = f"""# Network Consultant AI Report

## Report Details
- **Generated:** {datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S UTC')}
- **Request ID:** {data.get('request_id', 'N/A')}
- **Priority:** {data.get('priority', 'N/A')}
- **Confidence:** {data.get('confidence', 0) * 100:.1f}%

{'⚠️ **RED FLAG:** This issue requires immediate attention!' if data.get('red_flagged') else ''}

## Client Issue
{data.get('client_issue', 'N/A')}

## Consensus Analysis
{data.get('consensus', 'N/A')}

## Recommendations
{''.join(f'- {rec}\n' for rec in data.get('recommendations', []))}

## Agent Analysis
{''.join(f'### {name}\n{agent.get("analysis", "N/A")}\n\n' for name, agent in data.get('agents', {}).items())}
        """
        return md.encode('utf-8')

report_generator = ReportGenerator()